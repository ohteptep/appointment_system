from ast import Delete
from multiprocessing import context
from tkinter.messagebox import NO
from tkinter.tix import STATUS
from urllib import request
from django.shortcuts import redirect, render
from django.shortcuts import render
from django.http import HttpResponse

import json
from .models import *
from .forms import *
from django.contrib.auth import authenticate, login, logout
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.contrib.auth.views import PasswordChangeView
from django.urls import reverse_lazy
from django.core.mail import send_mail
from django.contrib.auth.hashers import check_password
from django.core import serializers
from django.http import JsonResponse
from django.utils.timezone import datetime

from datetime import date
from docx import Document
from docx2pdf import convert
import pythoncom
import os
import shutil

# Create your views here.
class PasswordsChangeView(PasswordChangeView):
    form_class = PasswordChangingForm
    success_url = reverse_lazy('login_student')
    
def index(request):
    return render(request, "index.html")

#LOGIN
def login_superuser(request):
    if request.user.is_authenticated:
        return redirect ('dashboard')
    else:
        if request.method == 'POST':
            userrr = request.POST.get('username')
            passw = request.POST.get('password') 
            user = authenticate(request, username=userrr,password=passw)
            if user is not None:
                get_superuser = depts.objects.get(username=userrr)
                if get_superuser.is_superuser == False:
                    messages.info(request,'Something went wrong')
                    return redirect('login_superuser')
                else:
                    login(request, user)
                    return redirect('dashboard')
            else:
                messages.info(request,'Username/Password is Incorrect')

    return render(request, "login_superuser.html")

def login_admin(request):
    if request.user.is_authenticated:
        return redirect ('admin_site')
    else:
        if request.method == 'POST':
            userrr = request.POST.get('username')
            passw = request.POST.get('password') 
            dept_names = request.POST.get('department')
            user = authenticate(request, username=userrr,password=passw)
            if user is not None:
                get_dept = depts.objects.get(username=userrr)
                emails = get_dept.email
                departs = get_dept.department
                firsts = get_dept.first_name
                lasts = get_dept.last_name
                pos = get_dept.position
                print(get_dept)
                if get_dept.department == "":
                    messages.info(request,'Something went wrong')
                    return redirect('login_admin')
                elif dept_names != get_dept.department:
                    messages.info(request,'Your Account is not in the right Department')
                    return redirect('login_admin')
                else:
                    login(request, user)
                    if dept_names == 'SD':
                        request.session['department'] = departs
                        return redirect('admin_site_sg')
                    elif dept_names == 'RE':
                        request.session['department'] = departs
                        return redirect('admin_site_re')
                    else:
                        request.session['email'] = emails
                        request.session['username'] = userrr
                        request.session['department'] = departs
                        request.session['first_name'] = firsts
                        request.session['last_name'] = lasts
                        request.session['position'] = pos
                        return redirect('admin_site')

            else:
                messages.info(request,'Username/Password is Incorrect')
    
    context = {
    }
    return render(request, "login_admin.html", context)

def login_student(request):
    if request.user.is_authenticated:
        return redirect ('book_app_student')
    else:
        if request.method == 'POST':
            userrr = request.POST.get('username')
            passw = request.POST.get('password') 
            user = authenticate(request, username=userrr,password=passw)
            if user is not None:
                get_staff = depts.objects.get(username=userrr)
                if get_staff.is_staff == False:
                    messages.info(request,'Something went wrong')
                    return redirect('login_student')
                else:
                    login(request, user)
                    request.session['username_student'] = userrr
                    return redirect('book_app_student')

            else:
                messages.info(request,'Username/Password is Incorrect')
    
    context = {
    }
    return render(request, "login_student.html", context)

#SIGNUP
def signup_student(request):
    signup = student_reg()
    if request.method == 'POST':
        signup = student_reg(request.POST)
        if signup.is_valid():
            signup.instance.is_staff = True
            signup.save()
            return redirect('login_student')
        else:
            messages.info(request,'Invalid Credentials!')

    context = {
        'signup': signup,
    }
   
    return render(request, "signup_student.html", context)

#APPOINTMENTS
def book_app(request):
    get_data = depts.objects.filter(is_staff = 0, is_active = 1).values()
    get_appointment = appointmentGuest(request.POST or None)
    if request.method == 'POST':
        if get_appointment.is_valid():
            print(get_appointment)
            get_appointment.save()
            messages.info(request,'Successfully Submitted')
            return redirect('book_app')
        else:
            messages.info(request,'Error Occured Submitting the Appointment Form')
    context = {
        'names': get_data
    }
    return render(request, "book_app.html", context)

@login_required(login_url='login_student')
def book_app_student(request):
    today = date.today()
    get_user = request.session['username_student']
    get_form_user = appointmentForm.objects.filter(username = get_user)
    get_data = depts.objects.filter(is_staff = 0, is_active = 1).values()
    get_appointment = appointmentGuest(request.POST or None)
    store_form_user_data = []
    for x in get_form_user:
        store_form_user_data.append(x)

    if request.method == 'POST':
        if get_appointment.is_valid():
            get_appointment.save()
            messages.info(request,'Successfully Submitted')
            return redirect('book_app_student')
        else:
            messages.info(request,'Error Occured Submitting the Appointment Form')
            return redirect('book_app_student')

    context = {
        'names': get_data,
        'username': get_user,
        'get_user_data': store_form_user_data,
        'date_today': today
    }
    return render(request, "book_app_student.html", context)


@login_required(login_url='login_student')
def book_app_alumni(request):
    get_user = request.session['username_student']
    get_form_user = appointmentForm.objects.filter(username = get_user)
    get_data = depts.objects.filter(is_staff = 0, is_active = 1).values()
    get_appointment = appointmentGuest(request.POST or None)

    store_form_user_data = []
    for x in get_form_user:
        store_form_user_data.append(x)

    if request.method == 'POST':
        if get_appointment.is_valid():
            get_appointment.save()
            messages.info(request,'Successfully Submitted')
            return redirect('book_app_alumni')
        else:
            messages.info(request,'Error Occured Submitting the Appointment Form')
            return redirect('book_app_alumni')

    context = {
        'names': get_data,
        'username': get_user,
        'get_user_data': store_form_user_data
    }
    return render(request, "book_app_alumni.html", context)

def css_form(request):
    get_css_form = formcss(request.POST or None)
    print(get_css_form)
    if request.method == 'POST':
        if get_css_form.is_valid():
            print(get_css_form)
            get_css_form.save()
            messages.info(request,'Successfully Submitted')
            return redirect('css_form')
        else:
            messages.info(request,'Error Occured Submitting the Appointment Form')
            return redirect('css_form')

    return render(request, "css_form.html")


#LOGOUT 
def logoutSuperuser(request):
    logout(request)
    return redirect('login_superuser')

def logoutAdmin(request):
    logout(request)
    return redirect('login_admin')

def logoutStudent(request):
    logout(request)
    return redirect('login_student')


#ADMIN
@login_required(login_url='login_admin')
def admin_site(request):
    get_dept = request.session['department']
    get_email_dept = request.session['email']
    get_first = request.session['first_name']
    get_last = request.session['last_name']
    get_position = request.session['position']

    if get_dept == "OAA":
        dept_names = 'Office of Academic Affair'
        check_sessions = get_position +' - '+ get_first +' ' + get_last +' - ' + dept_names + '('+get_dept+')'
    elif get_dept == "DIT":
        dept_names = 'Department of Information Technology'
        check_sessions = get_position +' - '+ get_first +' ' + get_last +' - ' + dept_names + '('+get_dept+')'
    elif get_dept == "DLA":
        dept_names = 'Department of Liberal Arts'
        check_sessions = get_position +' - '+ get_first +' ' + get_last +' - ' + dept_names + '('+get_dept+')'
    elif get_dept == "OCL":
        dept_names = 'Office of Campus Library'
        check_sessions = get_position +' - '+ get_first +' ' + get_last +' - ' + dept_names + '('+get_dept+')'
    elif get_dept == "DED":
        dept_names = 'Department of Education'
        check_sessions = get_position +' - '+ get_first +' ' + get_last +' - ' + dept_names + '('+get_dept+')'
    elif get_dept == "DMS":
        dept_names = 'Department of Mathematics and Science'
        check_sessions = get_position +' - '+ get_first +' ' + get_last +' - ' + dept_names + '('+get_dept+')'
    elif get_dept == "DOE":
        dept_names = 'Department of Engineering'
        check_sessions = get_position +' - '+ get_first +' ' + get_last +' - ' + dept_names + '('+get_dept+')'
    elif get_dept == "OSA":
        dept_names = 'Office of Student Affairs'
        check_sessions = get_position +' - '+ get_first +' ' + get_last +' - ' + dept_names + '('+get_dept+')'
    elif get_dept == "UITC":
        dept_names = 'University Information Technology Center '
        check_sessions = get_position +' - '+ get_first +' ' + get_last +' - ' + dept_names + '('+get_dept+')'
    elif get_dept == "DPE":
        dept_names = 'Department of Physical Education'
        check_sessions = get_position +' - '+ get_first +' ' + get_last +' - ' + dept_names + '('+get_dept+')'


    print(check_sessions)
    get_id_accept = request.POST.get('id_accept')
    get_id_declined = request.POST.get('id_decline')
    get_id_canceled = request.POST.get('id_cancel')
    get_id_delete = request.POST.get('id_delete')
    get_id_reapproved = request.POST.get('id_reapproved')

    get_name = request.POST.get('student_name')
    if get_name != None:
        composed_name_header = 'Good day,' + ' ' + get_name
        get_email = request.POST.get('student_email')
        hostemail = 'tupcappointment2022@gmail.com'
        msg = 'We would like you to fill-up our CSS form we already accepted your appointment.' + '\n' + 'Please click the link provided to open the css form' + ' http://127.0.0.1:8000/css_form/' + '\n' + 'Thank you have a nice day.' + '\n \n' + '- TUPC_APPOINTMENT_2022'  
        send_mail(
            composed_name_header,
            msg,
            hostemail,
            [get_email],
        )
        messages.info(request,'Css Form has been sent')

    
    if get_id_delete != None:
        delete_app = appointmentForm.objects.filter(id=get_id_delete)
        delete_app.delete()
        messages.info(request,'Successfully Deleted!')

    checkapp1 = appointmentForm.objects.filter(id = get_id_accept).update(status='APPROVED')
    if checkapp1 == 1:
        get_name = request.POST.get('student_name_accept')
        get_date_time = request.POST.get('dtime')
        get_departs = request.POST.get('depts')
        get_email_department = request.POST.get('department_email')
        appointmentForm.objects.filter(id = get_id_accept).update(notes='Your Appointment Successfully Approved')
        appointmentForm.objects.filter(id = get_id_accept).update(contactperson_email=get_email_department)
        appointmentForm.objects.filter(id = get_id_accept).update(sd_status='ONGOING')
        composed_name_header = 'Good day,' + ' ' + get_name
        get_email = request.POST.get('accept_email')
        hostemail = 'tupcappointment2022@gmail.com'


        msg = 'Your Appointment Successfully Approved' + '\n \n' + 'I would like to confirm your appointment with ' + get_departs + ' at ' + get_date_time + '\n \n' + 'Your appointment is scheduled to be held at TUPC ' + get_departs + '\n \n' + 'Please feel free to contact ' + get_departs + ' if you have any questions. We look forward to seeing you have a wonderful day! ' + '\n \n' + 'Regards,' + '\n' +  get_departs
        send_mail(
            composed_name_header,
            msg,
            hostemail,
            [get_email],
        )
        messages.info(request,'Successfully Approved')

    checkapp2 = appointmentForm.objects.filter(id = get_id_declined).update(status='DECLINED')
    if checkapp2 == 1:
        decline_compose = request.POST.get('decline_msg')
        get_name = request.POST.get('student_name_decline')
        appointmentForm.objects.filter(id = get_id_declined).update(notes=decline_compose)
        composed_name_header = 'Good day,' + ' ' + get_name
        get_email = request.POST.get('decline_email')
        hostemail = 'tupcappointment2022@gmail.com'
        msg = decline_compose + '\n \n' + '- TUPC_APPOINTMENT_2022'
        send_mail(
            composed_name_header,
            msg,
            hostemail,
            [get_email],
        )
        messages.info(request,'Successfully Declined')

    checkapp3 = appointmentForm.objects.filter(id = get_id_reapproved).update(status='APPROVED')
    if checkapp3 == 1:
        get_name = request.POST.get('student_name_reapprove')
        get_date_time_re = request.POST.get('dtime')
        get_departs_re = request.POST.get('depts')
        appointmentForm.objects.filter(id = get_id_reapproved).update(notes='Your Appointment Successfully Re-Approved')
        appointmentForm.objects.filter(id = get_id_reapproved).update(sd_status='ONGOING')
        composed_name_header = 'Good day,' + ' ' + get_name
        get_email = request.POST.get('reapprove_email')
        hostemail = 'tupcappointment2022@gmail.com'
        msg = 'Your Appointment Successfully Approved' + '\n \n' + 'We are confirming and re-approved the cancelled appointment that you have on ' + get_date_time_re + '. I apologize for the incovenience and short notice we hope you understand. ' + '\n \n' + 'We look forward to seeing you on ' + get_date_time_re + '. Have a wonderful day!' + '\n \n' + 'Regards,' + '\n' +  get_departs_re
        send_mail(
            composed_name_header,
            msg,
            hostemail,
            [get_email],
        )
        messages.info(request,'Successfully Approved')

    checkapp4= appointmentForm.objects.filter(id = get_id_canceled).update(status='DECLINED')
    if checkapp4 == 1:
        cancel_compose = request.POST.get('cancel_msg')
        get_name = request.POST.get('student_name_cancel')
        appointmentForm.objects.filter(id = get_id_canceled).update(notes=cancel_compose)
        composed_name_header = 'Good day,' + ' ' + get_name
        get_email = request.POST.get('cancel_email')
        hostemail = 'tupcappointment2022@gmail.com'
        msg = cancel_compose + '\n \n' + '- TUPC_APPOINTMENT_2022'
        send_mail(
            composed_name_header,
            msg,
            hostemail,
            [get_email],
        )
        messages.info(request,'Successfully Canceled')

    get_appointment_pending = appointmentForm.objects.filter(dept = get_dept, contactperson=check_sessions).filter(status='PENDING').values()
    get_appointment_approved = appointmentForm.objects.filter(dept = get_dept, contactperson=check_sessions).filter(status='APPROVED').values()
    get_appointment_declined = appointmentForm.objects.filter(dept = get_dept, contactperson=check_sessions).filter(status='DECLINED').values()
    get_appointment_history = appointmentForm.objects.filter(dept = get_dept, contactperson=check_sessions).values()

    # check_app = appointmentForm.objects.values()
    # for x in check_app:
    #     trys = x['contactperson']
    #     spliting = trys.split()
    #     for x in spliting:
    #         if get_last in x:
    #             print(x)
                
    if get_dept == "OAA":
        set_val = 'Office of Academic Affair'
        set_email = get_email_dept
    elif get_dept == "DIT":
        set_val = 'Department of Information Technology'
        set_email = get_email_dept
    elif get_dept == "DLA":
        set_val = 'Department of Liberal Arts'
        set_email = get_email_dept
    elif get_dept == "OCL":
        set_val = 'Office of Campus Library'
        set_email = get_email_dept
    elif get_dept == "DED":
        set_val = 'Department of Education'
        set_email = get_email_dept
    elif get_dept == "DMS":
        set_val = 'Department of Mathematics and Science'
        set_email = get_email_dept
    elif get_dept == "DOE":
        set_val = 'Department of Engineering'
        set_email = get_email_dept
    elif get_dept == "OSA":
        set_val = 'Office of Student Affairs'
        set_email = get_email_dept
    elif get_dept == "UITC":
        set_val = 'University Information Technology Center '
        set_email = get_email_dept
    elif get_dept == "DPE":
        set_val = 'Department of Physical Education'
        set_email = get_email_dept

    context = {
        'dept_name' : set_val,
        'dept_email': set_email,
        'dept_val_1': get_appointment_pending,
        'dept_val_2': get_appointment_approved, 
        'dept_val_3': get_appointment_declined, 
        'dept_val_4': get_appointment_history, 
    }


    return render(request, "admin_site.html", context)

@login_required(login_url='login_admin')
def admin_site_sg(request):
    today = date.today()
    print(today)
    get_appointment_approved = appointmentForm.objects.filter(status='APPROVED', pdate = today, sd_status='ONGOING').values()
    get_id_delete = request.POST.get('id_delete')
    get_email_check = request.POST.get('notify_email')

    if get_id_delete != None:
        appointmentForm.objects.filter(id=get_id_delete).update(sd_status='DONE')
        messages.info(request,'Successfully Deleted!')


    if get_email_check != None:
        get_message = request.POST.get('messages')
        composed_name_header = 'Student Appointment'
        get_email = request.POST.get('notify_email')
        hostemail = 'tupcappointment2022@gmail.com'
        msg = get_message + '\n \n' + '- TUPC_APPOINTMENT_2022'  
        send_mail(
            composed_name_header,
            msg,
            hostemail,
            [get_email],
        )
        messages.info(request,'Contact Perosonnel Successfully Notify')

    get_dept = request.session['department']
    if get_dept == 'SD':
        set_val = 'Security Department'

    context={
        'dept_name': set_val,
        'dept_val_1': get_appointment_approved
    }
    return render(request, "admin_site_sg.html", context)


@login_required(login_url='login_admin')
def admin_site_re(request):
    get_cssform = cssform.objects.all()
    get_dept = request.session['department']
    if get_dept == 'RE':
        set_val = 'Research & Extension'

    context={
        'dept_name': set_val,
        'cssform': get_cssform
    }
    return render(request, "admin_site_re.html", context)

#SUPERUSER
@login_required(login_url='login_superuser')
def dashboard(request):
    get_faculty = depts.objects.filter(is_staff = 0, is_superuser = 0).values()
    get_student = depts.objects.filter(is_staff = 1, is_superuser = 0).values()
    get_appointments = appointmentForm.objects.all()

    store_length_1 = len(get_faculty) 
    save_length_1 = [store_length_1]

    store_length_2 = len(get_student) 
    save_length_2 = [store_length_2]

    store_length_3 = len(get_appointments) 
    save_length_3 = [store_length_3]

    context = {
        'faculty': get_faculty,
        'length1': save_length_1,
        'length2': save_length_2,
        'length3': save_length_3,
    }
    return render(request, "dashboard.html", context)

@login_required(login_url='login_superuser')
def create_manage(request):
    get_faculty = depts.objects.filter(is_staff = 0).values()
    get_student = depts.objects.filter(is_staff = 1).filter(is_superuser = 0).values()
    get_app = appointmentForm.objects.all()

    get_id_update_admin = request.POST.get('id_update_admin')
    if get_id_update_admin != None:
        get_username = request.POST.get('e_username')
        get_first_name = request.POST.get('e_first_name')
        get_last_name = request.POST.get('e_last_name')
        get_email = request.POST.get('e_email')
        if depts.objects.filter(username = get_username, first_name= get_first_name, last_name= get_last_name, email=get_email).exists():
            messages.info(request,'No Changes Detected')
        else:
            depts.objects.filter(id=get_id_update_admin).update(username=get_username)
            depts.objects.filter(id=get_id_update_admin).update(first_name=get_first_name)
            depts.objects.filter(id=get_id_update_admin).update(last_name=get_last_name)
            depts.objects.filter(id=get_id_update_admin).update(email=get_email)
            messages.info(request,'Successfully Updated')

    get_id_delete_admin = request.POST.get('id_delete_admin')
    if get_id_delete_admin != None:
        delete_admin = depts.objects.filter(id=get_id_delete_admin)
        delete_admin.delete()
        messages.info(request,'Successfully Deleted!')

    get_id_disable_admin = request.POST.get('id_disable_admin')
    if get_id_disable_admin != None:
        if depts.objects.filter(id=get_id_disable_admin, is_active=False).exists():
            messages.info(request,'Account is Already Disabled')
        else:
            depts.objects.filter(id=get_id_disable_admin).update(is_active=False)
            messages.info(request,'Successfully Disabled the Account!')

    get_id_enable_admin = request.POST.get('id_enable_admin')
    if get_id_enable_admin != None:
        if depts.objects.filter(id=get_id_enable_admin, is_active=True).exists():
            messages.info(request,'Account is Already Enabled')
        else:
            depts.objects.filter(id=get_id_enable_admin).update(is_active=True)
            messages.info(request,'Successfully Enabled the Account!')

    get_id_update_student = request.POST.get('id_update_student')
    if get_id_update_student != None:
        get_username = request.POST.get('s_username')
        get_email = request.POST.get('s_email')
        if depts.objects.filter(username = get_username, email=get_email).exists():
            messages.info(request,'No Changes Detected')
        else:
            depts.objects.filter(id=get_id_update_student).update(username=get_username)
            depts.objects.filter(id=get_id_update_student).update(email=get_email)
            messages.info(request,'Successfully Updated')

    get_id_delete_student = request.POST.get('id_delete_student')
    if get_id_delete_student != None:
        delete_student = depts.objects.filter(id=get_id_delete_student)
        delete_student.delete()
        messages.info(request,'Successfully Deleted!')

    signup_admin = admin_reg()
    if request.method == 'POST':
        signup_admin = admin_reg(request.POST)
        get_pos = request.POST.get('position')
        get_dept = request.POST.get('department')
        get_idnum = request.POST.get('username')
        get_mail = request.POST.get('email')

        if get_pos == 'Head':
            if depts.objects.filter(position = 'Head', department= get_dept).exists():
                messages.info(request, 'Head Department is already exist')
            elif depts.objects.filter(username = get_idnum).exists():
                messages.info(request, 'I.D Number Already Exists')
            elif depts.objects.filter(email=get_mail).exists():
                messages.info(request, 'Email Already Exists')
            else:
                signup_admin = admin_reg(request.POST)
                if signup_admin.is_valid():
                    signup_admin.save()
                    messages.info(request,'Successfully Created Admin Account')
                    return redirect('create_manage')

        else:
            if depts.objects.filter(username = get_idnum).exists():
                messages.info(request, 'I.D Number Already Exists')
            elif depts.objects.filter(email=get_mail).exists():
                messages.info(request, 'Email Already Exists')
            else:
                signup_admin = admin_reg(request.POST)
                if signup_admin.is_valid():
                    signup_admin.save()
                    messages.info(request,'Successfully Created Admin Account')
                    return redirect('create_manage')
        

    context = {
        'signup_admin': signup_admin,
        'faculty': get_faculty,
        'student': get_student,
        'get_date': get_app
    }
    return render(request, "create_manage.html", context)

@login_required(login_url='login_superuser')
def appointments(request):
    get_appointment_pending = appointmentForm.objects.filter(status='PENDING').values()
    get_appointment_approved = appointmentForm.objects.filter(status='APPROVED').values()
    get_appointment_declined = appointmentForm.objects.filter(status='DECLINED').values()
    get_appointment_history = appointmentForm.objects.all()
    get_css_form = cssform.objects.all()

    store_length_pending = len(get_appointment_pending)
    save_length_pending = [store_length_pending]

    store_length_approved = len(get_appointment_approved)
    save_length_approved = [store_length_approved]

    store_length_decline = len(get_appointment_declined)
    save_length_decline = [store_length_decline]

    store_length_cssform = len(get_css_form)
    save_length_cssform = [store_length_cssform]


    context = {
        'get_length_pending': save_length_pending,
        'get_length_approved': save_length_approved,
        'get_length_declined': save_length_decline,
        'get_length_cssform': save_length_cssform,
        'dept_val': get_appointment_history, 
        'css_form_data': get_css_form
    }
    return render(request, "appointments.html", context)

@login_required(login_url='login_superuser')
def user(request):
    superuser = depts.objects.filter(is_superuser = True).values()
    get_id_superuser = request.POST.get('id_superuser')

    if get_id_superuser is not None:
        get_username = request.POST.get('username')
        get_first_name = request.POST.get('first_name')
        get_last_name = request.POST.get('last_name')
        get_email = request.POST.get('email')
        depts.objects.filter(id=get_id_superuser).update(username=get_username)
        depts.objects.filter(id=get_id_superuser).update(first_name=get_first_name)
        depts.objects.filter(id=get_id_superuser).update(last_name=get_last_name)
        depts.objects.filter(id=get_id_superuser).update(email=get_email)
        messages.info(request,'Successfully Updated')

    context = {
        'get_superuser': superuser
    }
    return render(request, "user.html", context)


#GENERATE PDF
def generatePDF(request):
    pythoncom.CoInitialize()
    directory = os.getcwd()
    print('DIRECTORY: '+directory)
    today = date.today()
    css_data = cssform.objects.all()
    doc = Document()

    pd1 = PDFS.objects.all()
    FILENO = len(pd1) + 1

    doc.add_heading('SAVED CSS FORM: '+str(today), 0)
    doc.add_heading('SAVED FILE NO: '+str(FILENO), 1)
    doc.add_paragraph(" ")
    for i in range(len(css_data)):
        doc.add_paragraph("FULL NAME: "+css_data[i].name)
        doc.add_paragraph("CONTACT: "+css_data[i].contact)
        doc.add_paragraph("EMAIL: "+css_data[i].email)
        doc.add_paragraph("PURPOSE: "+css_data[i].transaction)
        doc.add_paragraph("FEEDBACK: "+css_data[i].feedback)
        doc.add_paragraph("COMMENT: "+css_data[i].comment)
        doc.add_paragraph("- - - - - - - - - - - - - - - - - -")

    
    
    doc.save('SAVED-FILE-'+str(FILENO)+'-'+str(today)+'.docx')
    convert('SAVED-FILE-'+str(FILENO)+'-'+str(today)+'.docx')

    s = PDFS(PDFSave = 'SAVED-FILE-'+str(FILENO)+'-'+str(today)+'.pdf')
    s.save()

    shutil.move('SAVED-FILE-'+str(FILENO)+'-'+str(today)+'.pdf', 'pdf_files')
    os.remove('SAVED-FILE-'+str(FILENO)+'-'+str(today)+'.docx')

    #os.startfile('SAVED-'+str(today)+'.pdf')
    os.startfile(directory+'\pdf_files'+'\SAVED-FILE-'+str(FILENO)+'-'+str(today)+'.pdf')

    messages.info(request,'Successfully Generated a PDF file')

    return redirect('admin_site_re')


def notif(request):
    get_dept = request.session['department']
    get_appointment_pending = appointmentForm.objects.filter(dept = get_dept).filter(status='PENDING').values()
    get_length = len(get_appointment_pending) 
    return JsonResponse({'data':get_length})


def sd_notif(request):
    today = date.today()
    get_appointment_approved = appointmentForm.objects.filter(status='APPROVED', pdate = today, sd_status='ONGOING').values()
    get_length = len(get_appointment_approved) 
    return JsonResponse({'sd':get_length})